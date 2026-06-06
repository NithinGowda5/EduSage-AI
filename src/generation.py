import os
import requests
# pyrefly: ignore [missing-import]
import docx
from typing import List, Optional, Any

# pyrefly: ignore [missing-import]
from langchain_cohere import ChatCohere
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.language_models.chat_models import SimpleChatModel
from langchain_core.messages import BaseMessage
from langchain_core.callbacks.manager import CallbackManagerForLLMRun

# ---------------------------------------------------
# OPENROUTER WRAPPER
# ---------------------------------------------------

class ChatOpenRouter(SimpleChatModel):
    model_name: str
    openrouter_api_key: str
    temperature: float = 0.1

    @property
    def _llm_type(self) -> str:
        return "openrouter"

    def _call(
        self,
        messages: List[BaseMessage],
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        headers = {
            "Authorization": f"Bearer {self.openrouter_api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:8000",
            "X-Title": "RAG Research Assistant",
        }
        
        api_messages = []
        for m in messages:
            if m.type == "human":
                role = "user"
            elif m.type == "ai":
                role = "assistant"
            elif m.type == "system":
                role = "system"
            else:
                role = "user"
            api_messages.append({"role": role, "content": m.content})

        payload = {
            "model": self.model_name,
            "messages": api_messages,
            "temperature": self.temperature,
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=payload
        )
        response.raise_for_status()
        res_json = response.json()
        
        if "choices" in res_json and len(res_json["choices"]) > 0:
            return res_json["choices"][0]["message"]["content"]
        else:
            raise ValueError(f"Unexpected response structure from OpenRouter: {res_json}")

# ---------------------------------------------------
# LLM CONFIGURATION
# ---------------------------------------------------

def get_llm(temperature: float = 0.1, model: Optional[str] = None):
    provider = os.environ.get("PREFERRED_PROVIDER", "cohere")
    if provider == "openrouter":
        openrouter_api_key = os.environ.get("OPENROUTER_API_KEY")
        if not openrouter_api_key:
            raise ValueError("OpenRouter API Key is not set in the environment.")
        selected_model = model or os.environ.get("OPENROUTER_MODEL", "meta-llama/llama-3-8b-instruct:free")
        return ChatOpenRouter(
            model_name=selected_model,
            openrouter_api_key=openrouter_api_key,
            temperature=temperature
        )
    else:
        cohere_api_key = os.environ.get("COHERE_API_KEY")
        if not cohere_api_key:
            raise ValueError("Cohere API Key is not set in the environment.")
        cohere_model = model or "command-r-08-2024"
        return ChatCohere(model=cohere_model, temperature=temperature)

# ---------------------------------------------------
# QA PROMPT
# ---------------------------------------------------

qa_prompt_template = PromptTemplate(
    input_variables=["context", "question"],
    template="""
You are a precise research assistant.

Answer the user's research question based ONLY on the provided excerpts from multiple documents.

Cite each document when using its information
(example: [DocName, Page X]).

Note any agreements or contradictions between sources.

Context:
{context}

Research Question:
{question}

Answer:
"""
)

# ---------------------------------------------------
# RESEARCH REPORT PROMPT
# ---------------------------------------------------

report_prompt_template = PromptTemplate(
    input_variables=["context", "question"],
    template="""
You are an expert academic researcher.

Generate a structured research report based ONLY on the provided excerpts from multiple documents.

The report must include:

1. Introduction
2. Findings by Theme
3. Cross-Document Analysis
4. Contradictions Found
5. Conclusion
6. References

For every claim, cite the document name and page number
(example: [SourceDoc.pdf, p. 4]).

Context:
{context}

Research Question:
{question}

Structured Report:
"""
)

# ---------------------------------------------------
# FORMAT DOCUMENTS
# ---------------------------------------------------

def format_docs_for_context(docs):
    formatted = []

    for d in docs:
        source = d.metadata.get("source_document", "Unknown")
        page = d.metadata.get("page", "Unknown")

        formatted.append(
            f"--- Document: {source}, Page: {page} ---\n"
            f"{d.page_content}\n"
        )

    return "\n".join(formatted)

# ---------------------------------------------------
# GENERATE SHORT ANSWER
# ---------------------------------------------------

def generate_answer(question: str, retrieved_docs: list):

    llm = get_llm()

    context_str = format_docs_for_context(retrieved_docs)

    chain = qa_prompt_template | llm | StrOutputParser()

    response = chain.invoke({
        "context": context_str,
        "question": question
    })

    return response

# ---------------------------------------------------
# GENERATE FULL RESEARCH REPORT
# ---------------------------------------------------

def generate_research_report(question: str, retrieved_docs: list):

    context_str = format_docs_for_context(retrieved_docs)

    # Use the unified get_llm function with temperature=0.2
    provider = os.environ.get("PREFERRED_PROVIDER", "cohere")
    if provider == "openrouter":
        llm_report = get_llm(temperature=0.2)
    else:
        llm_report = get_llm(temperature=0.2, model="command-r-plus-08-2024")

    chain = report_prompt_template | llm_report | StrOutputParser()

    response = chain.invoke({
        "context": context_str,
        "question": question
    })

    return response

# ---------------------------------------------------
# EXPORT TO WORD
# ---------------------------------------------------

def export_to_word(
    report_text: str,
    filename: str = "Research_Report.docx"
) -> str:

    # Create folder if not exists
    base_data_dir = "/data" if (os.path.exists("/data") or os.environ.get("SPACE_ID")) else "data"
    os.makedirs(base_data_dir, exist_ok=True)

    doc = docx.Document()

    doc.add_heading("Research Report", level=0)

    for line in report_text.split("\n"):

        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)

        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)

        else:
            doc.add_paragraph(line)

    export_path = os.path.join(base_data_dir, filename)

    doc.save(export_path)

    return export_path

# ---------------------------------------------------
# EXAMPLE USAGE
# ---------------------------------------------------

"""
from langchain.schema import Document

docs = [
    Document(
        page_content="AI improves productivity.",
        metadata={
            "source_document": "paper1.pdf",
            "page": 2
        }
    )
]

question = "How does AI improve productivity?"

report = generate_research_report(question, docs)

path = export_to_word(report)

print(report)
print(f"Saved to: {path}")
"""